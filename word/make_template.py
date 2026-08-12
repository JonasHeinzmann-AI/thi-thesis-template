"""Build the Word version of the thesis template.

WHY A SCRIPT AND NOT A CHECKED-IN .DOCX ALONE
A .docx is a zip of XML. Reviewing a change to one in a pull request is not
possible, so the document is generated from this file and both are committed.
Edit the script, re-run it, commit both.

WHAT THE SCRIPT SETS THAT WORD'S DEFAULTS GET WRONG
  page      A4, margins 30 mm left and 20 mm right, the THI convention
  body      12 pt, 1.5 line spacing, justified, first-line indent off
  headings  real Heading 1..3 styles, so the table of contents field works and
            the navigation pane is usable
  toc       inserted as a FIELD, not as static text. Word fills it on F9 or on
            "Update Table". A typed-out list of headings looks the same until
            the first edit and then silently goes stale.

    python word/make_template.py
"""
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

BODY_PT, TITLE_PT = 12, 20


def field(par, code):
    """Insert a Word field. Used for the table of contents and page numbers."""
    r = par.add_run()
    for kind, text in (("begin", None), (None, code), ("separate", None),
                       (None, ""), ("end", None)):
        if kind:
            e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), kind)
        else:
            e = OxmlElement("w:instrText"); e.set(qn("xml:space"), "preserve"); e.text = text
        r._r.append(e)


def para(doc, text="", style=None, size=None, bold=False, align=None, space_after=None):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text); run.bold = bold
        if size:
            run.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def main():
    doc = Document()

    s = doc.sections[0]
    s.page_width, s.page_height = Mm(210), Mm(297)
    s.left_margin, s.right_margin = Mm(30), Mm(20)
    s.top_margin, s.bottom_margin = Mm(25), Mm(25)

    n = doc.styles["Normal"]
    n.font.name = "Times New Roman"; n.font.size = Pt(BODY_PT)
    n.paragraph_format.line_spacing = 1.5
    n.paragraph_format.space_after = Pt(6)
    n.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # ---------------------------------------------------------------- title page
    para(doc, space_after=0)
    para(doc, "Technische Hochschule Ingolstadt", size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Faculty of Computer Science", size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_after=48)
    para(doc, "Your Thesis Title, Which May Run Across Several Lines",
         size=TITLE_PT, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)
    para(doc, "Bachelor thesis", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)
    for label, value in (("submitted by", "Your Name"),
                         ("Programme", "Your Programme (B. Sc.)"),
                         ("First examiner", "Prof. Dr. First Examiner"),
                         ("Second examiner", "Prof. Dr. Second Examiner"),
                         ("Issued on", "TODO"),
                         ("Submitted on", "TODO")):
        p = para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        p.add_run("%s: " % label).bold = False
        p.add_run(value).bold = True
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ------------------------------------------------- declaration of authorship
    # The reason this template exists. Wording follows the THI declaration
    # extended with the AI-usage paragraph. Confirm it with the Pruefungsamt
    # before submitting: the source slide was marked preliminary. It also omits
    # the "not presented it elsewhere" clause the older template carries.
    # The serial commas in the quoted paragraphs below are the university's own
    # wording, reproduced verbatim. Do not copy-edit them.
    doc.add_heading("Declaration of Authorship", level=1)
    para(doc, "Declaration in accordance with § 30 Abs. 4 Nr. 7 APO THI", size=10,
         space_after=18)
    para(doc, "I hereby declare that I have written this paper independently, have not "
              "used any sources or aids other than those indicated, and have followed the "
              "principles of good scientific practice.",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)
    para(doc, "All content generated with the support of AI has been identified in "
              "accordance with recognized scientific principles. Exceptions to this "
              "identification include orthographic or grammatical corrections, "
              "translations, and improvements to wording that do not alter the meaning. "
              "I am aware that AI-generated content does not guarantee quality. I "
              "therefore declare that I have used AI tools solely as aids, that I have "
              "critically reviewed the AI-generated content, and that my own cognitive "
              "and creative influence predominates in this work. I affirm that I have "
              "fully understood the content of my work and can independently defend it.",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=48)
    para(doc, "Ingolstadt, ______________________", space_after=4)
    para(doc, "                          (Date)", size=9, space_after=48)
    para(doc, "______________________________", space_after=4)
    para(doc, "(Signature)", size=9, space_after=0)
    para(doc, "Your Name")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---------------------------------------------------------------- front matter
    doc.add_heading("Abstract", level=1)
    para(doc, "", space_after=0)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_heading("Acknowledgements", level=1)
    para(doc, "", space_after=0)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_heading("Table of Contents", level=1)
    # A field, so Word rebuilds it. Right-click the table and choose "Update
    # Field", or press F9 with the cursor inside it. It is empty until then,
    # which is expected rather than a fault.
    field(doc.add_paragraph(), r'TOC \o "1-3" \h \z \u')
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_heading("List of Abbreviations", level=1)
    t = doc.add_table(rows=1, cols=2); t.style = "Table Grid"
    t.rows[0].cells[0].text = "CNN"
    t.rows[0].cells[1].text = "convolutional neural network"
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---------------------------------------------------------------- chapters
    for i, (title, subs) in enumerate([
        ("Introduction", ["Motivation", "Problem Statement", "Contributions",
                          "Structure of this Thesis"]),
        ("Background and Related Work", ["Fundamentals", "Related Work"]),
        ("Method and Experimental Setup", ["Method", "Data", "Experimental Setup",
                                           "Evaluation Metrics"]),
        ("Results", ["Main Results", "Ablations"]),
        ("Discussion", ["Interpretation", "Limitations"]),
        ("Conclusion", ["Summary", "Future Work"]),
    ], start=1):
        doc.add_heading("%d  %s" % (i, title), level=1)
        for j, sub in enumerate(subs, start=1):
            doc.add_heading("%d.%d  %s" % (i, j, sub), level=2)
            para(doc, "", space_after=0)
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_heading("References", level=1)
    para(doc, "Use Word's own bibliography tool, Zotero or Citavi. Set the style once "
              "under References and keep it.", size=10)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading("Appendix", level=1)
    para(doc, "", space_after=0)

    # ---------------------------------------------------------------- page numbers
    for sec in doc.sections:
        p = sec.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        field(p, "PAGE")

    out = "word/thesis-template.docx"
    doc.save(out)
    print("geschrieben: %s" % out)


if __name__ == "__main__":
    main()
